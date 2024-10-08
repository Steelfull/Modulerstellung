from flask import request, redirect, url_for, Blueprint, render_template, send_file, jsonify
from .models import Modul
from . import db
from docx import Document
import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from weasyprint import HTML

main = Blueprint('main', __name__)

@main.route('/')
def index():
    return render_template('index.html')

@main.route('/save', methods=['POST'])
def save_module():
    # Felder aus dem Formular erfassen
    module_name = request.form['module_name']
    module_code = request.form['module_code']  # Modulkürzel
    department = request.form['department']
    module_type = request.form['module_type']
    qualifications = request.form.get('qualifications', None)  # Standardwert auf None setzen
    semesters = request.form.getlist('semesters')  # Holen Sie sich alle ausgewählten Semester
    description = request.form['description']
    sws = request.form['sws']
    ects = request.form['ects']
    lecture_type = request.form['lecture_type']
    

    presence_time = request.form.get('presence_time', type=int)
    self_study_time = request.form.get('self_study_time', type=int)

    # Validierung der Zeiten
    if presence_time is not None and presence_time < 0:
        return jsonify(success=False, message="Präsenzzeit darf nicht negativ sein.")
    if self_study_time is not None and self_study_time < 0:
        return jsonify(success=False, message="Selbststudiumszeit darf nicht negativ sein.")
    if qualifications is None:
        return jsonify({'success': False, 'message': 'Bitte wählen Sie eine Qualifikation aus.'}), 400

    total_hours = (presence_time or 0) + (self_study_time or 0)  # Berechnung der Gesamtstunden
    prerequisites_formal = request.form['prerequisites_formal']
    prerequisites_content = request.form['prerequisites_content']
    curriculum = request.form['curriculum']
    language_of_instruction = request.form['language_of_instruction']
    exam_type = request.form['exam_type']
    exam_format_duration = request.form['exam_format_duration']
    exam_language = request.form['exam_language']
    requirements_for_credits = request.form['requirements_for_credits']
    module_responsible = request.form['module_responsible']
    registration = request.form['registration']
    knowledge = request.form['knowledge']
    skills = request.form['skills']
    competences = request.form['competences']
    contents = request.form['contents']
    teaching_mode = request.form['teaching_mode']
    learning_mode = request.form['learning_mode']
    literature = request.form['literature']
    equipment_costs = request.form['equipment_costs']
    miscellaneous = request.form['miscellaneous']
    last_update = request.form['last_update']

    # Neues Modul speichern
    new_module = Modul(
        module_name=module_name,
        module_code=module_code,  # Modulkürzel speichern
        department=department,
        module_type=module_type,
        qualifications=qualifications,
        semesters=', '.join(semesters),
        description=description,
        sws=sws,
        ects=ects,
        lecture_type=lecture_type,
        presence_time=presence_time,
        self_study_time=self_study_time,
        total_hours=total_hours,  # Gesamtstunden speichern
        prerequisites_formal=prerequisites_formal,
        prerequisites_content=prerequisites_content,
        curriculum=curriculum,
        language_of_instruction=language_of_instruction,
        exam_type=exam_type,
        exam_format_duration=exam_format_duration,
        exam_language=exam_language,
        requirements_for_credits=requirements_for_credits,
        module_responsible=module_responsible,
        registration=registration,
        knowledge=knowledge,
        skills=skills,
        competences=competences,
        contents=contents,
        teaching_mode=teaching_mode,
        learning_mode=learning_mode,
        literature=literature,
        equipment_costs=equipment_costs,
        miscellaneous=miscellaneous,
        last_update=last_update
    )

    try:
        db.session.add(new_module)
        db.session.commit()
        return jsonify(success=True, message="Modul erfolgreich gespeichert.")
    except Exception as e:
        db.session.rollback()  # Rollback bei Fehler
        return jsonify(success=False, message="Fehler beim Speichern des Moduls: " + str(e))


@main.route('/check_hours', methods=['POST'])
def check_hours():
    # Lese die Daten aus der Anfrage
    presence_time = request.form.get('presence_time', type=int) or 0
    self_study_time = request.form.get('self_study_time', type=int) or 0
    total_hours = presence_time + self_study_time

    # Gebe die Gesamtstunden zurück als JSON-Antwort
    return jsonify(total_hours=total_hours)

@main.route('/modules')
def list_modules():
    # Hole alle gespeicherten Module aus der Datenbank
    modules = Modul.query.all()
    return render_template('modules.html', modules=modules)

@main.route('/export/<int:id>')
def export_module(id):
    module = Modul.query.get_or_404(id)
    filepath = os.path.join(os.path.dirname(__file__), f'{module.module_name}.docx')

    # Erstelle ein neues Word-Dokument
    doc = Document()

    # Titel und Modulname
    doc.add_heading('Modulbeschreibung', level=1)
    doc.add_heading(f'{module.module_name} ({module.module_code})', level=2)

    # Fachbereich
    doc.add_heading('Fachbereich / Abteilung:', level=3)
    doc.add_paragraph(module.department)

    # Kurzbeschreibung
    doc.add_heading('Kurzbeschreibung:', level=3)
    doc.add_paragraph(module.description)

    # SWS, ECTS, Arbeitsaufwand
    doc.add_heading('Studienumfang:', level=3)
    doc.add_paragraph(f'Semesterwochenstunden (SWS): {module.sws}')
    doc.add_paragraph(f'ECTS-Leistungspunkte (CP): {module.ects}')

    # Arbeitsaufwand (Zeitstunden)
    doc.add_heading('Arbeitsaufwand (Zeitstunden)', level=3)
    
    # Art der Lehrveranstaltungen
    doc.add_heading('Art der Lehrveranstaltungen:', level=4)
    doc.add_paragraph(f'{module.lecture_type} (z.B. Vorlesung, Seminar, Labor, Tutorium, …)')

    # Präsenz- und Selbststudium
    doc.add_heading('Präsenz (Zeitstunden):', level=4)
    doc.add_paragraph(str(module.presence_time))

    doc.add_heading('Selbststudium (Zeitstunden):', level=4)
    doc.add_paragraph(str(module.self_study_time))

    doc.add_heading('Gesamt (Zeitstunden):', level=4)
    doc.add_paragraph(str(module.total_hours))  # Gesamtstunden anzeigen

    # Überfachliche Qualifikationen
    doc.add_heading('Überfachliche Qualifikationen:', level=3)
    doc.add_paragraph(f'Ist auch in anderen Studiengängen relevant: {"Ja" if module.cross_program else "Nein"}')

    # Voraussetzungen
    doc.add_heading('Voraussetzungen:', level=3)
    doc.add_paragraph(f'Formal: {module.prerequisites_formal}')
    doc.add_paragraph(f'Inhaltlich: {module.prerequisites_content}')

    # Zuordnung zum Curriculum
    doc.add_heading('Zuordnung zum Curriculum:', level=3)
    doc.add_paragraph(f'{module.curriculum}')

    # Unterrichtssprache
    doc.add_heading('Unterrichtssprache:', level=3)
    doc.add_paragraph(f'{module.language_of_instruction}')

    # Prüfungsart
    doc.add_heading('Prüfungsart:', level=3)
    doc.add_paragraph(f'{module.exam_type}')

    # Prüfungsform, -dauer, -umfang
    doc.add_heading('Prüfungsform, -dauer, -umfang:', level=3)
    doc.add_paragraph(f'{module.exam_format_duration}')

    # Prüfungsbedingungen
    doc.add_heading('Prüfungssprache:', level=3)
    doc.add_paragraph(f'{module.exam_language}')

    # Voraussetzungen zum Erwerb der Leistungspunkte
    doc.add_heading('Voraussetzungen zum Erwerb der Leistungspunkte:', level=3)
    doc.add_paragraph(module.requirements_for_credits)

    # Modulverantwortliche*r
    doc.add_heading('Modulverantwortliche*r:', level=3)
    doc.add_paragraph(f'{module.module_responsible}')

    # Anmeldung über
    doc.add_heading('Anmeldung über:', level=3)
    doc.add_paragraph(f'{module.registration}')

    # Lernergebnisse und Kompetenzen
    doc.add_heading('Lernergebnisse und Kompetenzen:', level=3)
    doc.add_paragraph('Kenntnisse:', style='Heading4')
    doc.add_paragraph(f'{module.knowledge}')
    doc.add_paragraph('Fertigkeiten:', style='Heading4')
    doc.add_paragraph(f'{module.skills}')
    doc.add_paragraph('Kompetenzen:', style='Heading4')
    doc.add_paragraph(f'{module.competences}')

    # Inhalte
    doc.add_heading('Inhalte:', level=3)
    doc.add_paragraph(f'{module.contents}')

    # Lehrmodus und Lernmodus
    doc.add_heading('Lehrmodus:', level=3)
    doc.add_paragraph(f'{module.teaching_mode}')
    doc.add_heading('Lernmodus:', level=3)
    doc.add_paragraph(f'{module.learning_mode}')

    # Literatur
    doc.add_heading('Literatur:', level=3)
    doc.add_paragraph(f'{module.literature}')

    # Ausstattungskosten
    doc.add_heading('Ausstattungskosten:', level=3)
    doc.add_paragraph(f'{module.equipment_costs}')

    # Sonstiges
    doc.add_heading('Sonstiges:', level=3)
    doc.add_paragraph(f'{module.miscellaneous}')

    # Letzte Änderung
    doc.add_heading('Letzte Änderung:', level=3)
    doc.add_paragraph(f'{module.last_update}')

    # Speichere das Dokument
    doc.save(filepath)

    return send_file(filepath, as_attachment=True)

@main.route('/export/pdf/<int:id>')
def export_module_pdf(id):
    module = Modul.query.get_or_404(id)
    filepath = os.path.join(os.path.dirname(__file__), f'{module.module_name}.pdf')

    # Erstelle ein PDF-Dokument
    c = canvas.Canvas(filepath, pagesize=letter)
    width, height = letter

    # Titel und Modulname
    c.setFont("Helvetica-Bold", 20)
    c.drawString(72, height - 72, 'Modulbeschreibung')
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, height - 100, f'{module.module_name} ({module.module_code})')

    # Fachbereich
    c.setFont("Helvetica", 12)
    c.drawString(72, height - 140, 'Fachbereich / Abteilung:')
    c.drawString(72, height - 160, module.department)

    # Kurzbeschreibung
    c.drawString(72, height - 180, 'Kurzbeschreibung:')
    c.drawString(72, height - 200, module.description)

    # SWS, ECTS, Arbeitsaufwand
    c.drawString(72, height - 220, 'Studienumfang:')
    c.drawString(72, height - 240, f'Semesterwochenstunden (SWS): {module.sws}')
    c.drawString(72, height - 260, f'ECTS-Leistungspunkte (CP): {module.ects}')

    # Arbeitsaufwand (Zeitstunden)
    c.drawString(72, height - 280, 'Arbeitsaufwand (Zeitstunden)')
    
    # Art der Lehrveranstaltungen
    c.drawString(72, height - 300, 'Art der Lehrveranstaltungen:')
    c.drawString(72, height - 320, f'{module.lecture_type}')

    # Präsenz- und Selbststudium
    c.drawString(72, height - 340, 'Präsenz (Zeitstunden):')
    c.drawString(72, height - 360, str(module.presence_time))

    c.drawString(72, height - 380, 'Selbststudium (Zeitstunden):')
    c.drawString(72, height - 400, str(module.self_study_time))

    c.drawString(72, height - 420, 'Gesamt (Zeitstunden):')
    c.drawString(72, height - 440, str(module.total_hours))

    # Überfachliche Qualifikationen
    c.drawString(72, height - 460, 'Überfachliche Qualifikationen:')
    c.drawString(72, height - 480, f'Ist auch in anderen Studiengängen relevant: {"Ja" if module.cross_program else "Nein"}')

    # Voraussetzungen
    c.drawString(72, height - 500, 'Voraussetzungen:')
    c.drawString(72, height - 520, f'Formal: {module.prerequisites_formal}')
    c.drawString(72, height - 540, f'Inhaltlich: {module.prerequisites_content}')

    # Zuordnung zum Curriculum
    c.drawString(72, height - 560, 'Zuordnung zum Curriculum:')
    c.drawString(72, height - 580, f'{module.curriculum}')

    # Unterrichtssprache
    c.drawString(72, height - 600, 'Unterrichtssprache:')
    c.drawString(72, height - 620, f'{module.language_of_instruction}')

    # Prüfungsart
    c.drawString(72, height - 640, 'Prüfungsart:')
    c.drawString(72, height - 660, f'{module.exam_type}')

    # Prüfungsform, -dauer, -umfang
    c.drawString(72, height - 680, 'Prüfungsform, -dauer, -umfang:')
    c.drawString(72, height - 700, f'{module.exam_format_duration}')

    # Prüfungsbedingungen
    c.drawString(72, height - 720, 'Prüfungssprache:')
    c.drawString(72, height - 740, f'{module.exam_language}')

    # Voraussetzungen zum Erwerb der Leistungspunkte
    c.drawString(72, height - 760, 'Voraussetzungen zum Erwerb der Leistungspunkte:')
    c.drawString(72, height - 780, module.requirements_for_credits)

    # Modulverantwortliche*r
    c.drawString(72, height - 800, 'Modulverantwortliche*r:')
    c.drawString(72, height - 820, f'{module.module_responsible}')

    # Anmeldung über
    c.drawString(72, height - 840, 'Anmeldung über:')
    c.drawString(72, height - 860, f'{module.registration}')

    # Lernergebnisse und Kompetenzen
    c.drawString(72, height - 880, 'Lernergebnisse und Kompetenzen:')
    c.drawString(72, height - 900, 'Kenntnisse:')
    c.drawString(72, height - 920, f'{module.knowledge}')
    c.drawString(72, height - 940, 'Fertigkeiten:')
    c.drawString(72, height - 960, f'{module.skills}')
    c.drawString(72, height - 980, 'Kompetenzen:')
    c.drawString(72, height - 1000, f'{module.competences}')

    # Inhalte
    c.drawString(72, height - 1020, 'Inhalte:')
    c.drawString(72, height - 1040, f'{module.contents}')

    # Lehrmodus und Lernmodus
    c.drawString(72, height - 1060, 'Lehrmodus:')
    c.drawString(72, height - 1080, f'{module.teaching_mode}')
    c.drawString(72, height - 1100, 'Lernmodus:')
    c.drawString(72, height - 1120, f'{module.learning_mode}')

    # Literatur
    c.drawString(72, height - 1140, 'Literatur:')
    c.drawString(72, height - 1160, f'{module.literature}')

    # Ausstattungskosten
    c.drawString(72, height - 1180, 'Ausstattungskosten:')
    c.drawString(72, height - 1200, f'{module.equipment_costs}')

    # Sonstiges
    c.drawString(72, height - 1220, 'Sonstiges:')
    c.drawString(72, height - 1240, f'{module.miscellaneous}')

    # Letzte Änderung
    c.drawString(72, height - 1260, 'Letzte Änderung:')
    c.drawString(72, height - 1280, f'{module.last_update}')

    # Speichere das PDF-Dokument
    c.save()

    return send_file(filepath, as_attachment=True)

@main.route('/delete/<int:id>', methods=['POST'])
def delete_module(id):
    module = Modul.query.get_or_404(id)  # Hole das Modul anhand der ID
    db.session.delete(module)              # Lösche das Modul
    db.session.commit()                    # Bestätige die Änderungen
    return jsonify(success=True, message="Modul erfolgreich gelöscht.")


@main.route('/export/html/<int:id>')
def export_module_html(id):
    module = Modul.query.get_or_404(id)
    
    # HTML-Inhalt generieren
    html_content = f"""
    <!DOCTYPE html>
    <html lang="de">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Modulbeschreibung - {module.module_name}</title>
    </head>
    <body>
        <h1>Modulbeschreibung: {module.module_name}</h1>
        <h2>Modulkürzel: {module.module_code}</h2>
        <p><strong>Fachbereich:</strong> {module.department}</p>
        <p><strong>Kurzbeschreibung:</strong> {module.description}</p>
        <p><strong>SWS:</strong> {module.sws}</p>
        <p><strong>ECTS:</strong> {module.ects}</p>
        <p><strong>Art der Lehrveranstaltungen:</strong> {module.lecture_type}</p>
        <p><strong>Präsenzzeit:</strong> {module.presence_time} Stunden</p>
        <p><strong>Selbststudium:</strong> {module.self_study_time} Stunden</p>
        <p><strong>Gesamtstunden:</strong> {module.total_hours} Stunden</p>
        <p><strong>Modultyp:</strong> {module.module_type} </p>
        <p><strong>Überfachliche Qualifikationen:</strong> {module.qualifications}</p>
        <p><strong>Wird angeboten im:</strong> {module.semesters}</p>
        <p><strong>Voraussetzungen (formal):</strong> {module.prerequisites_formal}</p>
        <p><strong>Voraussetzungen (inhaltlich):</strong> {module.prerequisites_content}</p>
        <p><strong>Zuordnung zum Curriculum:</strong> {module.curriculum}</p>
        <p><strong>Unterrichtssprache:</strong> {module.language_of_instruction}</p>
        <p><strong>Prüfungsart:</strong> {module.exam_type}</p>
        <p><strong>Prüfungsform, -dauer, -umfang:</strong> {module.exam_format_duration}</p>
        <p><strong>Prüfungssprache:</strong> {module.exam_language}</p>
        <p><strong>Voraussetzungen für Leistungspunkte:</strong> {module.requirements_for_credits}</p>
        <p><strong>Modulverantwortliche*r:</strong> {module.module_responsible}</p>
        <p><strong>Anmeldung über:</strong> {module.registration}</p>
        <p><strong>Kentnisse:</strong> {module.knowledge}</p>
        <p><strong>Fertigkeiten:</strong> {module.skills}</p>
        <p><strong>Kompetenzen:</strong> {module.competences}</p>
        <p><strong>Inhalte:</strong> {module.contents}</p>
        <p><strong>Lehrmodus:</strong> {module.teaching_mode}</p>
        <p><strong>Lernmodus:</strong> {module.learning_mode}</p>
        <p><strong>Literatur:</strong> {module.literature}</p>
        <p><strong>Ausstattungskosten:</strong> {module.equipment_costs}</p>
        <p><strong>Sonstiges:</strong> {module.miscellaneous}</p>
        <p><strong>Letzte Aktualisierung:</strong> {module.last_update}</p>
    </body>
    </html>
    """

    # Speichere den HTML-Inhalt in einer Datei
    html_filename = os.path.join(os.path.dirname(__file__), f'{module.module_name}.html')

    with open(html_filename, 'w', encoding='utf-8') as file:
        file.write(html_content)

    # Sende die HTML-Datei als Download
    return send_file(html_filename, as_attachment=True)



@main.route('/export/pdf_from_html/<int:id>')
def export_pdf_from_html(id):
    module = Modul.query.get_or_404(id)

    # HTML-Inhalt generieren
    html_content = render_template('module_template.html', module=module)  # Verwenden Sie eine HTML-Vorlage

    # Erstelle eine PDF-Datei mit WeasyPrint
    pdf_filename = os.path.join(os.path.dirname(__file__), f'{module.module_name}.pdf')
    HTML(string=html_content).write_pdf(pdf_filename, stylesheets=['app\static\export_styles.css'])  # Fügen Sie den Stylesheet hinzu

    # Sende die PDF-Datei als Download
    return send_file(pdf_filename, as_attachment=True)
